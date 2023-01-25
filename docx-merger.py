from docx import Document
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, Cm
from docxcompose.composer import Composer
from docx.enum.table import WD_ALIGN_VERTICAL
import os, sys, json

# file_time = "1674028907.json"
file_time = sys.argv[1]
with open(file_time, "r") as file:
    msg_received = json.load(file)

FILES = "file"
SAVE_FOLDER = "saved files"
if not os.path.exists(SAVE_FOLDER):
    os.mkdir(SAVE_FOLDER)
if not os.path.exists(FILES):
    os.mkdir(FILES)


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(run):
    fldChar1 = create_element("w:fldChar")
    create_attribute(fldChar1, "w:fldCharType", "begin")

    instrText = create_element("w:instrText")
    create_attribute(instrText, "xml:space", "preserve")
    instrText.text = "PAGE"

    fldChar2 = create_element("w:fldChar")
    create_attribute(fldChar2, "w:fldCharType", "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


document = Document("sample.docx")
section = document.sections[0]
# section.top_margin = Inches(0.39)
# section.bottom_margin = Inches(0.39)
# section.left_margin = Inches(0.39)
# section.right_margin = Inches(0.39)
section.header_distance = Inches(0.48)
section.footer_distance = Inches(0.51)
header = section.header
header.is_linked_to_previous = True
paragraph = header.paragraphs[0]
paragraph.text = "<header>"

font_styles = document.styles
font_charstyle = font_styles.add_style("CommentsStyle", WD_STYLE_TYPE.CHARACTER)
font = font_charstyle.font
font.name = "Calibri"
font.size = Pt(16)
font.bold = True

document.add_paragraph(msg_received["output"])


p = document.add_paragraph()
p.add_run("Table of Contents", style="CommentsStyle")

paragraph = document.add_paragraph()

run = paragraph.add_run()

fldChar = OxmlElement("w:fldChar")
fldChar.set(qn("w:fldCharType"), "begin")
fldChar.set(qn("w:dirty"), "true")

instrText = OxmlElement("w:instrText")
instrText.set(qn("xml:space"), "preserve")
instrText.text = 'TOC \o "1-3" \h \z \\u'

fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "separate")

fldChar5 = OxmlElement("w:updateFileds")
fldChar5.set(qn("w:val"), "true")

fldChar4 = OxmlElement("w:fldChar")
fldChar4.set(qn("w:fldCharType"), "end")

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
r_element.append(fldChar5)
document.add_page_break()
p_element = paragraph._p

document.add_paragraph(msg_received["output"])
# data = ((1, "Reasoning"), (2, "Maths"), (3, "MCQs"))  # Add Dynamic Data
table = document.add_table(rows=1, cols=2, style="Table Grid")

table.columns[1].width = Inches(18)
row = table.rows[0].cells
row[0].text = "Sr No"
row[1].text = "Subject"
row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
make_rows_bold(table.rows[0])
i = 1
for name in msg_received["subject"]:
    row = table.add_row().cells
    row[0].text = str(i)
    row[1].text = name
    row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    i = i + 1

for i in range(len(table.rows)):
    table.rows[i].height = Inches(0.4)

document.add_paragraph("\n")

document.add_paragraph(msg_received["desc"])
files = msg_received["code"]  # Files Array
# files = ["A6-M-H.docx", "A1040-M-E.docx"]  # Files Array

composer = Composer(document)
for file in files:
    document.add_page_break()
    f = FILES + "/" + file + ".docx"
    doc1 = Document(f)
    composer.append(doc1)

add_page_number(document.sections[0].footer.paragraphs[0].add_run())
document.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

name = SAVE_FOLDER + "/" + msg_received["output"]
document.save(name + ".docx")
